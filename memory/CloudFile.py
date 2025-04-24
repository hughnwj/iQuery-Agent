# ==================================================
# 导入依赖区域
# ==================================================

import os
import shutil
import tempfile

from docx import Document

# 基础路径设置
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
base_path = os.path.join(BASE_DIR, "data", "doc")


# ==================================================
# CloudFile 类区域
# ==================================================

class CloudFile():
    """
    云文档操作类，用于管理项目文档的创建、读取、更新和删除等操作
    """

    def __init__(self, project_name, part_name, doc_content=None):
        """
        初始化云文档对象
        :param project_name: 项目名称，即项目文件夹名称
        :param part_name: 项目某部分名称，即项目文件名称
        :param doc_content: 项目文件初始内容，默认为None
        """
        self.project_name = project_name
        self.part_name = part_name
        folder_path = create_or_get_folder(folder_name=project_name)
        self.doc_list = list_files_in_folder(folder_name=project_name)
        file_path = create_or_get_doc(folder_name=self.project_name,
                                      doc_name=self.part_name)
        self.doc_content = doc_content
        if doc_content != None:
            append_content_in_doc(folder_name=project_name,
                                  doc_name=part_name,
                                  qa_string=doc_content)

    def get_doc_content(self):
        """
        获取文档内容
        :return: 返回文档内容字符串
        """
        self.doc_content = get_file_content(folder_name=self.project_name, doc_name=self.part_name)
        return self.doc_content

    def append_doc_content(self, content):
        """
        向文档追加内容
        :param content: 要追加的内容，可以是字典列表或消息对象列表
        """
        formatted_content = []
        for msg in content:
            if isinstance(msg, dict):
                role = msg.get('role', 'unknown')
                msg_content = msg.get('content', '')
                formatted_content.append(f"{role}: {msg_content}")
            else:
                try:
                    formatted_content.append(f"{msg.role}: {msg.content}")
                except AttributeError:
                    formatted_content.append(str(msg))
        content_str = "\n".join(formatted_content)
        append_content_in_doc(folder_name=self.project_name,
                              doc_name=self.part_name,
                              qa_string=content_str)

    def clear_content(self):
        """
        清空文档内容
        """
        clear_content_in_doc(folder_name=self.project_name, doc_name=self.part_name)

    def delete_all_files(self):
        """
        删除项目文件夹内所有文件
        """
        delete_all_files_in_folder(folder_name=self.project_name)

    def update_doc_list(self):
        """
        更新文档列表
        """
        self.doc_list = list_files_in_folder(folder_name=self.project_name)

    def rename_doc(self, new_name):
        """
        重命名文档
        :param new_name: 新文档名称
        :return: 返回新文档名称
        """
        self.part_name = rename_doc(folder_name=self.project_name,
                                    doc_name=self.part_name,
                                    new_name=new_name)


# ==================================================
# 文件操作函数区域
# ==================================================

def create_or_get_folder(folder_name):
    """
    创建或获取文件夹路径
    :param folder_name: 文件夹名称
    :return: 返回完整文件夹路径
    """
    full_path = os.path.join(base_path, folder_name)
    if not os.path.exists(full_path):
        os.makedirs(full_path)
        print(f"目录 {folder_name} 创建成功")
    else:
        print(f"目录 {folder_name} 已存在")


def create_or_get_doc(folder_name, doc_name):
    """
    创建或获取文档路径
    :param folder_name: 文件夹名称
    :param doc_name: 文档名称
    :return: 返回完整文档路径
    """
    full_path_folder = os.path.join(base_path, folder_name)
    file_path_doc = os.path.join(base_path + "/" + folder_name, f'{doc_name}.doc')
    if not os.path.exists(full_path_folder):
        os.makedirs(full_path_folder)
    if os.path.exists(file_path_doc):
        document = Document(file_path_doc)
    else:
        document = Document()
    document.save(file_path_doc)
    return file_path_doc


def append_content_in_doc(folder_name, doc_name, qa_string):
    """
    向文档追加内容
    :param folder_name: 文件夹名称
    :param doc_name: 文档名称
    :param qa_string: 要追加的内容字符串
    """
    full_path_folder = base_path + "/" + folder_name
    full_path_doc = os.path.join(full_path_folder, doc_name) + ".doc"
    if not os.path.exists(full_path_folder):
        os.makedirs(full_path_folder)
    if os.path.exists(full_path_doc):
        document = Document(full_path_doc)
    else:
        document = Document()
    print(qa_string)
    document.add_paragraph(qa_string)
    document.save(full_path_doc)
    print(f"内容已追加到 {doc_name}")


def get_file_content(folder_name, doc_name):
    """
    获取文件内容
    :param folder_name: 文件夹名称
    :param doc_name: 文档名称
    :return: 返回文档内容字符串或错误信息
    """
    file_path = os.path.join(folder_name, doc_name)
    full_path = os.path.join(base_path, file_path) + ".doc"
    if not os.path.exists(full_path):
        return "文件不存在"
    try:
        doc = Document(full_path)
        content = []
        for para in doc.paragraphs:
            content.append(para.text)
        return '\n'.join(content)
    except Exception as e:
        return f"读取文件时发生错误: {e}"


def clear_content_in_doc(folder_name, doc_name):
    """
    清空文档内容
    :param folder_name: 文件夹名称
    :param doc_name: 文档名称
    """
    file_path = os.path.join(base_path + "/" + folder_name, f'{doc_name}.doc')
    doc = Document(file_path)
    for p in doc.paragraphs:
        for run in p.runs:
            run.text = ''
    doc.save(file_path)
    print("文档内容清除完毕")


def list_files_in_folder(folder_name):
    """
    列出文件夹内所有文件
    :param folder_name: 文件夹名称
    :return: 返回文件名列表
    """
    full_path = os.path.join(base_path, folder_name)
    file_names = [f for f in os.listdir(full_path) if os.path.isfile(os.path.join(full_path, f))]
    return file_names


def rename_doc(folder_name, doc_name, new_name):
    """
    重命名文档
    :param folder_name: 文件夹名称
    :param doc_name: 原文档名称
    :param new_name: 新文档名称
    :return: 返回新文档名称
    """
    file_path = os.path.join(base_path + "/" + folder_name, f'{doc_name}.doc')
    new_file_path = os.path.join(base_path + "/" + folder_name, f'{new_name}.doc')
    os.rename(file_path, new_file_path)
    return new_name


def delete_all_files_in_folder(folder_name):
    """
    删除文件夹内所有文件
    :param folder_name: 文件夹名称
    """
    full_path = os.path.join(base_path, folder_name)
    for filename in os.listdir(full_path):
        file_path = os.path.join(full_path, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
                print("文件已清除完毕")
        except Exception as e:
            print('Failed to delete %s. Reason: %s' % (file_path, e))


def append_img_in_doc(folder_name, doc_name, fig):
    """
    向文档追加图片
    :param folder_name: 文件夹名称
    :param doc_name: 文档名称
    :param fig: matplotlib的Figure对象
    """
    full_path_folder = base_path + "/" + folder_name
    full_path_doc = os.path.join(full_path_folder, doc_name) + ".doc"
    if not os.path.exists(full_path_folder):
        os.makedirs(full_path_folder)
    if os.path.exists(full_path_doc):
        print(full_path_doc)
        document = Document(full_path_doc)
    else:
        document = Document()
    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
        fig.savefig(tmpfile.name, format='png')
        document.add_picture(tmpfile.name)
    document.save(full_path_doc)
    print(f"图片已追加到 {doc_name}")
